import tempfile
import unittest
from io import BytesIO
from pathlib import Path

from docx import Document
from fastapi import UploadFile

from app.services.document_parser import extract_text_from_docx
from app.services.file_store import save_upload_to_temp


ALLOWED_DOCX_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _build_docx_bytes() -> bytes:
    buffer = BytesIO()
    doc = Document()
    doc.add_paragraph("Hello world")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Table text"
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _make_upload(content: bytes, filename: str = "sample.docx", content_type: str = ALLOWED_DOCX_TYPE) -> UploadFile:
    return UploadFile(filename=filename, file=BytesIO(content), headers={"content-type": content_type})


class FileStoreTests(unittest.TestCase):
    def test_save_valid_docx_to_temp(self):
        upload = _make_upload(_build_docx_bytes())
        path = save_upload_to_temp(upload)
        self.addCleanup(lambda: path.unlink(missing_ok=True))

        self.assertTrue(path.exists())
        self.assertEqual(path.suffix, ".docx")
        self.assertGreater(path.stat().st_size, 0)

    def test_rejects_invalid_extension(self):
        upload = _make_upload(b"payload", filename="notes.txt")
        with self.assertRaises(ValueError):
            save_upload_to_temp(upload)

    def test_rejects_invalid_content_type(self):
        upload = _make_upload(_build_docx_bytes(), content_type="text/plain")
        with self.assertRaises(ValueError):
            save_upload_to_temp(upload)

    def test_enforces_max_size(self):
        upload = _make_upload(b"a" * 20, content_type=ALLOWED_DOCX_TYPE)
        with self.assertRaises(ValueError):
            save_upload_to_temp(upload, max_bytes=10)


class DocumentParserTests(unittest.TestCase):
    def test_extracts_paragraphs_and_tables(self):
        data = _build_docx_bytes()
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(data)
            tmp_path = Path(tmp.name)
        self.addCleanup(lambda: tmp_path.unlink(missing_ok=True))

        extracted = extract_text_from_docx(tmp_path)
        self.assertIn("Hello world", extracted)
        self.assertIn("Table text", extracted)

    def test_rejects_wrong_extension_path(self):
        with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as tmp:
            tmp.write(b"not a docx")
            tmp_path = Path(tmp.name)
        self.addCleanup(lambda: tmp_path.unlink(missing_ok=True))

        with self.assertRaises(ValueError):
            extract_text_from_docx(tmp_path)

    def test_rejects_corrupted_docx(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(b"not a real docx file")
            tmp_path = Path(tmp.name)
        self.addCleanup(lambda: tmp_path.unlink(missing_ok=True))

        with self.assertRaises(ValueError):
            extract_text_from_docx(tmp_path)


if __name__ == "__main__":
    unittest.main()