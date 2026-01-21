"""Minimal smoke checks for local FastAPI server."""
from __future__ import annotations

import tempfile
from pathlib import Path

from docx import Document
import requests

BASE_URL = "http://127.0.0.1:8000"


def test_health():
    r = requests.get(f"{BASE_URL}/health")
    r.raise_for_status()
    data = r.json()
    assert data.get("status") == "ok", data
    print("/health", data)


def test_root():
    r = requests.get(BASE_URL)
    r.raise_for_status()
    print("/", r.status_code, "bytes:", len(r.text))


def test_rewrite_download():
    doc = Document()
    doc.add_heading("Sample Agreement", level=2)
    doc.add_paragraph("This agreement is made between the parties on the date signed.")
    doc.add_paragraph("The parties agree to the following terms.")

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp_path = Path(tmp.name)
    doc.save(tmp_path)

    try:
        with tmp_path.open("rb") as handle:
            files = {
                "file": (
                    "sample.docx",
                    handle,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            }
            r = requests.post(f"{BASE_URL}/api/rewrite", files=files)
            r.raise_for_status()
            output_path = Path("rewritten_output.docx")
            output_path.write_bytes(r.content)
            print("rewrite ->", output_path.resolve())
    finally:
        tmp_path.unlink(missing_ok=True)


if __name__ == "__main__":
    test_health()
    test_root()
    test_rewrite_download()
