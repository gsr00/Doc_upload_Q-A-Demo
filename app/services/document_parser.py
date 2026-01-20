"""DOCX parsing utilities."""
from __future__ import annotations

from pathlib import Path
from zipfile import BadZipFile

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

ALLOWED_EXTENSIONS: set[str] = {".docx"}


def _validate_path(path: Path | None) -> Path:
    if path is None:
        raise ValueError("File path is required.")
    if not path.exists():
        raise ValueError("File not found.")
    if path.suffix.lower() not in ALLOWED_EXTENSIONS:
        raise ValueError("Unsupported file type; only .docx is supported.")
    return path


def extract_text_from_docx(path: Path) -> str:
    """Extract plain text from a DOCX file.

    Returns a single string with paragraphs separated by newlines. Raises
    ValueError when the file is missing, of the wrong type, or corrupted.
    """
    _validate_path(path)
    try:
        doc = Document(path)
    except (PackageNotFoundError, BadZipFile, OSError, ValueError) as exc:
        raise ValueError("Invalid or corrupted DOCX file.") from exc

    parts: list[str] = []

    # Paragraphs outside of tables
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Capture text inside tables as well
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if text:
                        parts.append(text)

    return "\n".join(parts)