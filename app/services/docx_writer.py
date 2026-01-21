"""DOCX writer utilities for rewritten text."""
from __future__ import annotations

from pathlib import Path

from docx import Document


def _is_heading(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    if stripped.isupper() and len(stripped) <= 80:
        return True
    if stripped.endswith(":") and len(stripped) <= 100:
        return True
    return False


def write_docx(text: str, output_path: Path) -> Path:
    """Write rewritten text to a DOCX file with simple headings/paragraphs."""
    doc = Document()
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if _is_heading(stripped):
            doc.add_heading(stripped.rstrip(":"), level=2)
        else:
            doc.add_paragraph(stripped)
    doc.save(output_path)
    return output_path
